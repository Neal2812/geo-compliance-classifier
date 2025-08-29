"""DOCX parser using python-docx."""

from pathlib import Path
from typing import List

from ..logging_conf import get_logger
from ..schema import DocumentArtifact

logger = get_logger(__name__)


def parse_docx(file_path: Path) -> DocumentArtifact:
    """Parse DOCX file and extract text.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        DocumentArtifact with extracted text
    """
    doc_id = file_path.stem
    warnings = []
    text = ""
    
    try:
        from docx import Document
        
        doc = Document(str(file_path))
        paragraphs = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        text = "\n".join(paragraphs)
        logger.debug(f"Extracted {len(text)} characters from {file_path}")
        
    except ImportError:
        warnings.append("python-docx not available")
    except Exception as e:
        warnings.append(f"Failed to parse DOCX: {e}")
        logger.warning(f"Failed to parse {file_path}: {e}")
    
    if not text:
        warnings.append("No text could be extracted from DOCX")
    
    return DocumentArtifact(
        doc_id=doc_id,
        doc_type="docx",
        source_path=str(file_path),
        raw_text=text,
        parse_warnings=warnings
    )
